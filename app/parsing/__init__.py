from __future__ import annotations

import csv
from enum import Enum
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pydantic import BaseModel, Field
from pypdf import PdfReader


class DocumentType(str, Enum):
    PDF = "pdf"
    WORD = "word"
    SPREADSHEET = "spreadsheet"
    CSV = "csv"


class DocumentLocation(BaseModel):
    """A source fragment and its position within an uploaded document."""

    kind: str
    text: str = ""
    page: int | None = None
    index: int | None = None
    table_index: int | None = None
    row: int | None = None
    column: int | None = None
    sheet: str | None = None


class ParsedDocument(BaseModel):
    document_type: DocumentType
    text: str = ""
    locations: list[DocumentLocation] = Field(default_factory=list)


class ParseDocumentError(Exception):
    """Raised when a supported document cannot be read."""


class UnsupportedDocumentType(ParseDocumentError):
    """Raised when a document extension is not supported."""


def parse_document(path: Path) -> ParsedDocument:
    """Extract text and source locations from a supported local document."""

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path)
    if suffix == ".docx":
        return _parse_word(path)
    if suffix in {".xlsx", ".xlsm"}:
        return _parse_spreadsheet(path)
    if suffix == ".csv":
        return _parse_csv(path)
    raise UnsupportedDocumentType(f"Unsupported document type: {suffix or 'no extension'}")


def _parse_pdf(path: Path) -> ParsedDocument:
    try:
        reader = PdfReader(path)
        locations = []
        fragments = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            locations.append(DocumentLocation(kind="page", page=page_number, text=text))
            if text:
                fragments.append(text)
        return ParsedDocument(
            document_type=DocumentType.PDF,
            text="\n".join(fragments),
            locations=locations,
        )
    except Exception as error:
        raise ParseDocumentError("Unable to parse PDF document.") from error


def _parse_word(path: Path) -> ParsedDocument:
    try:
        document = Document(path)
        locations = []
        fragments = []
        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if text:
                locations.append(DocumentLocation(kind="paragraph", index=index, text=text))
                fragments.append(text)
        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                for column_index, cell in enumerate(row.cells, start=1):
                    text = cell.text.strip()
                    if text:
                        locations.append(
                            DocumentLocation(
                                kind="table_cell",
                                table_index=table_index,
                                row=row_index,
                                column=column_index,
                                text=text,
                            )
                        )
                        fragments.append(text)
        return ParsedDocument(
            document_type=DocumentType.WORD,
            text="\n".join(fragments),
            locations=locations,
        )
    except Exception as error:
        raise ParseDocumentError("Unable to parse Word document.") from error


def _parse_spreadsheet(path: Path) -> ParsedDocument:
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
        locations = []
        fragments = []
        for worksheet in workbook.worksheets:
            fragments.append(f"Sheet: {worksheet.title}")
            for row_index, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
                values = [str(value) for value in row if value is not None and str(value) != ""]
                if values:
                    text = " | ".join(values)
                    locations.append(
                        DocumentLocation(
                            kind="spreadsheet_row",
                            sheet=worksheet.title,
                            row=row_index,
                            text=text,
                        )
                    )
                    fragments.append(text)
        workbook.close()
        return ParsedDocument(
            document_type=DocumentType.SPREADSHEET,
            text="\n".join(fragments),
            locations=locations,
        )
    except Exception as error:
        raise ParseDocumentError("Unable to parse spreadsheet document.") from error


def _parse_csv(path: Path) -> ParsedDocument:
    try:
        locations = []
        fragments = []
        with path.open("r", encoding="utf-8-sig", newline="") as source:
            for row_index, row in enumerate(csv.reader(source), start=1):
                values = [value for value in row if value]
                if values:
                    text = " | ".join(values)
                    locations.append(DocumentLocation(kind="csv_row", row=row_index, text=text))
                    fragments.append(text)
        return ParsedDocument(
            document_type=DocumentType.CSV,
            text="\n".join(fragments),
            locations=locations,
        )
    except Exception as error:
        raise ParseDocumentError("Unable to parse CSV document.") from error


__all__ = [
    "DocumentLocation",
    "DocumentType",
    "ParseDocumentError",
    "ParsedDocument",
    "UnsupportedDocumentType",
    "parse_document",
]
