from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from app.parsing import ParseDocumentError, UnsupportedDocumentType, parse_document


def test_parses_word_paragraphs_and_table_cells_with_locations(tmp_path: Path) -> None:
    path = tmp_path / "notice.docx"
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Alice"
    document.save(path)

    parsed = parse_document(path)

    assert parsed.document_type == "word"
    assert "First paragraph" in parsed.text
    assert "Name" in parsed.text
    assert "Alice" in parsed.text
    assert any(location.kind == "paragraph" and location.index == 1 for location in parsed.locations)
    assert any(
        location.kind == "table_cell"
        and location.table_index == 1
        and location.row == 1
        and location.column == 2
        for location in parsed.locations
    )


def test_parses_excel_all_sheets_with_sheet_and_row_locations(tmp_path: Path) -> None:
    path = tmp_path / "workbook.xlsx"
    workbook = Workbook()
    summary = workbook.active
    summary.title = "Summary"
    summary.append(["Item", "Amount"])
    summary.append(["Apples", 12])
    details = workbook.create_sheet("Details")
    details.append(["Owner", "Status"])
    details.append(["Alice", "Open"])
    workbook.save(path)

    parsed = parse_document(path)

    assert parsed.document_type == "spreadsheet"
    assert "Item" in parsed.text
    assert "Apples" in parsed.text
    assert "Details" in parsed.text
    assert "Alice" in parsed.text
    assert any(location.sheet == "Summary" and location.row == 2 for location in parsed.locations)
    assert any(location.sheet == "Details" and location.row == 2 for location in parsed.locations)


def test_parses_utf8_bom_csv_with_header_and_row_locations(tmp_path: Path) -> None:
    path = tmp_path / "records.csv"
    path.write_text("Name,City\nAlice,Shanghai\n", encoding="utf-8-sig")

    parsed = parse_document(path)

    assert parsed.document_type == "csv"
    assert "Name" in parsed.text
    assert "Alice" in parsed.text
    assert any(location.row == 1 for location in parsed.locations)
    assert any(location.row == 2 for location in parsed.locations)


def test_empty_pdf_returns_empty_text_without_error(tmp_path: Path) -> None:
    path = tmp_path / "empty.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with path.open("wb") as output:
        writer.write(output)

    parsed = parse_document(path)

    assert parsed.document_type == "pdf"
    assert parsed.text == ""
    assert any(location.page == 1 for location in parsed.locations)


def test_rejects_unsupported_extension(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("not supported", encoding="utf-8")

    with pytest.raises(UnsupportedDocumentType, match="\.txt"):
        parse_document(path)


def test_wraps_unreadable_pdf_without_exposing_file_content(tmp_path: Path) -> None:
    path = tmp_path / "broken.pdf"
    secret = "private file contents must not appear"
    path.write_text(secret, encoding="utf-8")

    with pytest.raises(ParseDocumentError) as error:
        parse_document(path)

    assert secret not in str(error.value)
